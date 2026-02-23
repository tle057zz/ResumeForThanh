import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, size: int = 12, bold: bool = True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)


def find_first_nonempty_paragraph_indices(doc: Document, count: int = 3):
    idxs = []
    for i, p in enumerate(doc.paragraphs):
        if p.text and p.text.strip():
            idxs.append(i)
            if len(idxs) >= count:
                break
    return idxs


def insert_after(doc: Document, paragraph_index: int, text: str = ""):
    # Workaround: python-docx lacks direct insert API; we add at end and swap text when needed.
    # For simplicity, we just append new content at the end to avoid corrupting structure.
    return doc.add_paragraph(text)


def tailor(base_path: Path, out_path: Path):
    doc = Document(str(base_path))

    # Heuristic: capture top block (name/title), then append tailored sections after the document
    # to avoid disturbing original formatting. Recruiters will see tailored summary first in v2.

    # Add a page break-like separator only if document is empty at end (skip to keep compact)

    # Tailored Summary for Rokt
    add_heading(doc, "Summary (Rokt — Network Integrity)")
    summary = (
        "Curious, data-driven analyst with hands-on SQL and Python experience. Comfortable preparing,\n"
        "cleaning, and interpreting data; building simple dashboards; and collaborating with engineers\n"
        "and data scientists. Recently built a visitor-integrity web app (Flask + SQLite) that logs\n"
        "traffic signals (IP, user-agent, headers), distinguishes human vs automated activity, and\n"
        "exposes REST APIs with caching/compression — directly relevant to Rokt's Network Integrity."
    )
    p = doc.add_paragraph(summary)
    p.runs[0].font.size = Pt(10)

    # Rokt-Focused Highlights
    add_heading(doc, "Network Integrity Highlights")
    add_bullets(doc, [
        "Analyzed traffic patterns and recent-activity windows (e.g., 5‑minute online counts)",
        "Implemented heuristics from user-agent, headers, geo/IP to flag automated traffic",
        "Built endpoints for integrity metrics; documented behavior and deployment considerations",
        "Created visual components and charts to communicate trends to non-technical stakeholders",
    ])

    # Skills Focus
    add_heading(doc, "Skills Focus for Rokt")
    add_bullets(doc, [
        "SQL for joins, aggregations, window functions; spreadsheets for quick analyses",
        "Python for data prep/automation (pandas, requests, JSON)",
        "Dashboards/visualization (Chart.js workflow; Tableau/Power BI familiarity)",
        "Flask APIs, SQLite persistence, gzip/Brotli compression, cache headers",
        "Clear communication and collaboration with engineers, data scientists, and PMs",
    ])

    # Intent
    add_heading(doc, "Role Intent")
    intent = (
        "Excited to grow as a Junior Business Analyst on the Network Integrity team — contributing to\n"
        "data preparation, dashboards, and integrity signal analysis while learning from a collaborative\n"
        "team that values ownership and innovation."
    )
    p2 = doc.add_paragraph(intent)
    p2.runs[0].font.size = Pt(10)

    doc.save(str(out_path))


def main():
    parser = argparse.ArgumentParser(description="Tailor a CV for Rokt based on an existing base .docx")
    parser.add_argument("--base", required=False, default="docs/cv-resumes/Nho_Thanh_Le_BA_Simplus_CV.docx")
    parser.add_argument("--out", required=False, default="docs/cv-resumes/Nho_Thanh_Le_Rokt_Junior_BA_CV_v2.docx")
    args = parser.parse_args()

    base = Path(args.base).resolve()
    out = Path(args.out).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    tailor(base, out)
    print(f"Tailored CV written to: {out}")


if __name__ == "__main__":
    main()




