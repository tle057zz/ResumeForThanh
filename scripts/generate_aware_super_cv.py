from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, size: int = 12, bold: bool = True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)


def main() -> Path:
    out_path = Path(__file__).resolve().parents[1] / "docs" / "cv-resumes" / "Nho_Thanh_Le_Aware_Super_Tech_Internship_CV.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Header
    name_p = doc.add_paragraph()
    name_run = name_p.add_run("Nho Thanh Le")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Technology Summer Intern — Data & AI | Enterprise | Operations")
    title_run.font.size = Pt(11)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run("Sydney, Australia | 0432 813 909 | lenhothanh.nsl@gmail.com | linkedin.com/in/tle057/")
    contact_run.font.size = Pt(9)

    # Summary aligned to Aware Super
    add_heading(doc, "Summary")
    summary = (
        "Postgraduate student in Data Management & Analytics with hands‑on experience delivering production‑style projects. "
        "Curious learner with a growth mindset, eager to contribute across Technology Operations, Data & AI, and Enterprise Services. "
        "Built real projects with Python/Flask, SQLite, React/Vue, and implemented performance, accessibility, and reliability improvements."
    )
    p = doc.add_paragraph(summary)
    p.runs[0].font.size = Pt(10)

    # Availability & eligibility
    add_heading(doc, "Availability & Eligibility")
    add_bullets(doc, [
        "Available full‑time for 12 weeks from 24 Nov 2025 (Mon–Fri)",
        "Sydney location; hybrid work with 2+ days in‑office",
        "Valid Australian working rights during internship period",
        "Minimum Credit average; final‑year postgraduate student",
    ])

    # Skills
    add_heading(doc, "Skills")
    add_bullets(doc, [
        "Data & AI: Python (Pandas, NumPy), ETL, basic ML workflow (PyTorch/Transformers setup)",
        "Engineering: REST APIs (Flask), SQLite, JSON, testing mindset, Git",
        "Frontend: React, Vue, HTML/CSS/Bootstrap, accessibility awareness",
        "Cloud/Infra: Caching, compression (gzip/Brotli), basic CI/CD familiarity",
        "Communication: Documentation, stakeholder updates, collaborative problem solving",
    ])

    # Relevant Experience / Projects
    add_heading(doc, "Relevant Experience & Projects")
    add_bullets(doc, [
        "Portfolio Platform (Flask, SQLite, React/Vue): Built APIs for blog posts, analytics, and caching; ensured persistence and reliability.",
        "Performance Optimisation: Implemented gzip/Brotli, HTTP caching, lazy loading; improved first‑load time after deployments.",
        "Visitor Insights: Parsed headers/user‑agents/IPs to understand traffic integrity; added dashboards and endpoints.",
        "Accessibility & UX: Improved contrast, keyboard focus, and responsive behavior; used animations judiciously for readability.",
        "Algorithm Demonstrations: Exposed Pascal’s Triangle and other utilities to showcase problem solving and testing.",
    ])

    # Areas of Interest (match tracks)
    add_heading(doc, "Areas of Interest")
    add_bullets(doc, [
        "Technology Operations & Support: Service Desk exposure, asset tracking, platform hygiene",
        "Data & AI: Data engineering pipelines, model integration, responsible AI design",
        "Enterprise Services: Network & Cloud foundations, test automation practices",
    ])

    # Education
    add_heading(doc, "Education")
    add_bullets(doc, [
        "Master of IT & IT Management — University of Sydney (in progress)",
        "Bachelor of Science in IT — University of Technology Sydney (Distinction)",
    ])

    # Values fit
    add_heading(doc, "Values & Impact")
    add_bullets(doc, [
        "Motivated by member‑first outcomes and high‑trust teams",
        "Committed to inclusion, accessibility, and continuous learning",
        "Enjoy mentoring/being mentored; receptive to feedback and iteration",
    ])

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = main()
    print(f"Wrote CV to: {path}")



