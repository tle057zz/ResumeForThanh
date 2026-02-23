from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_heading(paragraph_text: str, document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run(paragraph_text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_bullets(items, document: Document) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10)


def main() -> Path:
    out_path = Path(__file__).resolve().parents[1] / "docs" / "cv-resumes" / "Nho_Thanh_Le_Rokt_Junior_BA_CV.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Name / Title
    name_p = doc.add_paragraph()
    name_run = name_p.add_run("Nho Thanh Le")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Junior Business Analyst — Network Integrity")
    title_run.font.size = Pt(11)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run("Email: lenhothanh.nsl@gmail.com  |  Portfolio: https://www.thanhle.it.com  |  Sydney, AU")
    contact_run.font.size = Pt(9)

    # Profile
    add_heading("Profile", doc)
    profile = (
        "Curious, data-driven analyst who enjoys finding signal in noisy traffic data. "
        "Hands-on with SQL and Python (Flask, pandas), comfortable building dashboards and "
        "instrumenting data pipelines. Recently engineered a visitor-tracking web app that "
        "detects human vs automated traffic, logs geo/IP metadata, and exposes REST APIs, "
        "which aligns directly with Rokt's Network Integrity mission."
    )
    p = doc.add_paragraph(profile)
    p.runs[0].font.size = Pt(10)

    # Core Skills
    add_heading("Core Skills", doc)
    add_bullets([
        "SQL for data preparation and exploration (joins, CTEs, window functions)",
        "Python for analysis and automation (pandas, requests, JSON, REST)",
        "Visualization and dashboards (Tableau/Power BI equivalent workflows, Chart.js)",
        "Web analytics and event tracking (Flask, endpoints, caching, compression)",
        "Traffic quality signals: user-agent, headers, geo-IP, recency/activity windows",
        "Clear communication with cross-functional partners (engineers, data scientists, PMs)",
    ], doc)

    # Relevant Projects (tailored)
    add_heading("Relevant Projects", doc)
    add_bullets([
        "Visitor Integrity Analytics App — Built a Flask backend recording visits, IP, user-agent, "
        "and page views in SQLite; implemented human vs automated heuristics; added REST APIs for "
        "stats and persistence; optimized with gzip/Brotli and aggressive caching.",
        "Career Advice Posts Persistence — Replaced localStorage with server-backed CRUD APIs, ensured durable "
        "storage in SQLite, added loading states and error handling.",
        "Interactive Portfolio — React/Vue animations, lazy-loading, and preload hints to improve first paint; "
        "deployed routes and asset paths for fast, reliable loads.",
    ], doc)

    # Experience (concise, outcome-focused)
    add_heading("Experience Highlights", doc)
    add_bullets([
        "Analyzed traffic patterns and created summaries for recent activity windows (5-min online count)",
        "Built metrics endpoints and visual components to make integrity signals easy to interpret",
        "Partnered with engineering-style tasks: API design, DB schema evolution, and caching strategies",
        "Documented persistence and deployment behavior; improved time-to-first-byte via compression",
    ], doc)

    # Education (based on provided context)
    add_heading("Education", doc)
    add_bullets([
        "Postgraduate Student in Data Management & Analytics (in progress)",
    ], doc)

    # Tools
    add_heading("Tools & Technologies", doc)
    add_bullets([
        "SQL, Python (pandas, requests), Flask, SQLite",
        "JavaScript, React.js, Vue.js, Bootstrap 5, Chart.js",
        "REST APIs, JSON, caching, gzip/Brotli (Flask-Compress)",
        "Git, GitHub, macOS, shell scripting",
    ], doc)

    # Close with intent
    add_heading("Role Fit", doc)
    fit = (
        "Motivated to grow as a Business Analyst within Network Integrity. Comfortable supporting data prep, "
        "building simple dashboards, and partnering with engineers and data scientists to improve fraud/IVT detection. "
        "Excited by Rokt's culture of ownership, learning, and AI-driven innovation."
    )
    p2 = doc.add_paragraph(fit)
    p2.runs[0].font.size = Pt(10)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = main()
    print(f"Wrote CV to: {path}")




