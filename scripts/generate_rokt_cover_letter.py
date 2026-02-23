from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main() -> Path:
    out_path = Path(__file__).resolve().parents[1] / "docs" / "cv-resumes" / "Nho_Thanh_Le_Rokt_Junior_BA_Cover_Letter.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Header
    name_p = doc.add_paragraph()
    name_run = name_p.add_run("Nho Thanh Le")
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run("Sydney, Australia | lenhothanh.nsl@gmail.com | www.thanhle.it.com")
    contact_run.font.size = Pt(9)

    doc.add_paragraph("")

    # Date and company
    doc.add_paragraph("Hiring Team")
    doc.add_paragraph("Rokt")
    doc.add_paragraph("Sydney, NSW")
    doc.add_paragraph("")

    # Greeting
    p = doc.add_paragraph("Dear Hiring Team,")
    p.runs[0].font.size = Pt(11)

    # Paragraph 1: Motivation and fit
    para1 = (
        "I am excited to apply for the Junior Business Analyst role on the Network Integrity team at Rokt. "
        "I am early in my career and highly motivated to grow in data analysis and ecommerce. My recent work "
        "building a visitor-integrity web application—where I instrumented traffic logging, distinguished "
        "human vs automated activity using headers and user-agent signals, and exposed REST APIs for analytics—"
        "aligns closely with the goals of Network Integrity."
    )
    p1 = doc.add_paragraph(para1)
    p1.runs[0].font.size = Pt(11)

    # Paragraph 2: Relevant skills
    para2 = (
        "I enjoy working with data and feel comfortable using SQL and Python to prepare, clean, and interpret "
        "datasets. I have experience creating simple dashboards and visualizations (e.g., Chart.js workflows) to "
        "communicate insights to both technical and nontechnical audiences. On the engineering side, I have built and "
        "maintained Flask-based APIs, implemented gzip/Brotli compression and cache headers to improve performance, and "
        "persisted analytics in SQLite with clear documentation for deployment and behavior."
    )
    p2 = doc.add_paragraph(para2)
    p2.runs[0].font.size = Pt(11)

    # Paragraph 3: Collaboration & learning
    para3 = (
        "I collaborate well with engineers, data scientists, and product stakeholders—asking questions, listening, and "
        "sharing clear summaries. I am detail‑oriented and conscientious, knowing that accurate data supports stronger "
        "decisions. I am particularly drawn to Rokt’s culture of ownership, transparency, and learning, and I’m excited "
        "by the opportunity to contribute to integrity systems that protect clients and improve data quality."
    )
    p3 = doc.add_paragraph(para3)
    p3.runs[0].font.size = Pt(11)

    # Paragraph 4: Close
    para4 = (
        "Thank you for your time and consideration. I would welcome the chance to discuss how my curiosity, hands‑on "
        "approach, and practical experience can support the Network Integrity team. I am looking forward to the next steps "
        "in the process."
    )
    p4 = doc.add_paragraph(para4)
    p4.runs[0].font.size = Pt(11)

    # Sign-off
    doc.add_paragraph("")
    p5 = doc.add_paragraph("Sincerely,")
    p5.runs[0].font.size = Pt(11)
    p6 = doc.add_paragraph("Nho Thanh Le")
    p6.runs[0].font.size = Pt(11)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = main()
    print(f"Wrote cover letter to: {path}")




